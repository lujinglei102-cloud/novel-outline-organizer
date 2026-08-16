# -*- coding: utf-8 -*-
"""生成《小说大纲梳理器》开发计划 v1.0"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if level == 0:
            run.font.size = Pt(18)
        elif level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.line_spacing = 1.0
    return h


def add_para(doc, text, bold=False, size=10.5, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.bold = bold
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_mono(doc, text, size=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        set_cell_bg(hdr_cells[i], 'D9D9D9')
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============== 文档 ==============
doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.0
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ============== 头 ==============
add_heading(doc, '开发计划', level=0)
add_table(doc,
    ['字段', '内容'],
    [
        ['标题', '小说大纲梳理器 开发计划'],
        ['版本', 'v1.0'],
        ['日期', '2026-08-16'],
        ['依据', '《小说大纲梳理器PRD_v1.0.docx》'],
        ['状态', '待评审'],
    ],
    col_widths=[3, 13])
add_para(doc, '')
add_para(doc, '变更记录：', bold=True, size=11)
add_table(doc,
    ['版本', '日期', '说明'],
    [['v1.0', '2026-08-16', '初版，基于 PRD v1.0 拆解到可执行任务级']],
    col_widths=[2, 3, 11])

doc.add_page_break()

# ============== 1. 技术栈与项目规范 ==============
add_heading(doc, '1. 技术栈与项目规范', level=1)

add_heading(doc, '1.1 技术选型', level=2)
add_table(doc,
    ['层级', '技术', '理由'],
    [
        ['框架', 'React 18 + TypeScript', '生态成熟，类型安全，组件化适合卡片墙'],
        ['构建', 'Vite 5', '开发热更新快，生产构建体积小'],
        ['路由', 'React Router v6', 'SPA 路由 /cards /sort /outline /export'],
        ['状态', 'Zustand', '轻量，避免 Redux 模板代码'],
        ['样式', 'Tailwind CSS 3', '响应式原子类，移动端适配快'],
        ['UI 组件', 'Headless UI + 自研', '弹窗/下拉用 Headless，卡片墙自研'],
        ['数据层', 'Dexie.js（IndexedDB 封装）', 'Promise API，本地持久化，预留同步接口'],
        ['图表', 'ECharts 5', '情绪曲线支持拖拽节点'],
        ['分词', 'jieba-wasm', '浏览器内中文分词，0 后端'],
        ['导出', 'file-saver + marked', 'MD/TXT 导出'],
        ['部署', 'Vercel / Netlify', '静态托管，0 服务器成本'],
    ],
    col_widths=[2.5, 4.5, 9])

add_heading(doc, '1.2 项目目录结构', level=2)
add_mono(doc,
    'src/\n'
    '  main.tsx                # 入口\n'
    '  App.tsx                 # 路由 + 全局布局\n'
    '  routes/\n'
    '    CardsPage.tsx         # /cards 卡片墙\n'
    '    SortPage.tsx          # /sort 梳理结果\n'
    '    OutlinePage.tsx       # /outline 骨架构建\n'
    '    ExportPage.tsx        # /export 大纲导出\n'
    '  components/             # 通用组件\n'
    '    Navbar.tsx\n'
    '    CardThumb.tsx\n'
    '    CardEditModal.tsx\n'
    '    EmotionChart.tsx\n'
    '    TimelineDivider.tsx\n'
    '    ...\n'
    '  stores/                 # Zustand 状态\n'
    '    cardStore.ts\n'
    '    sortStore.ts\n'
    '    outlineStore.ts\n'
    '  db/                     # Dexie 数据层\n'
    '    database.ts           # 表定义\n'
    '    cardRepo.ts           # 卡片 CRUD\n'
    '    templateRepo.ts\n'
    '  engine/                 # 纯规则梳理引擎\n'
    '    narrativeLine.ts      # 叙事线排序\n'
    '    characterExtract.ts   # 角色提取\n'
    '    foreshadowLink.ts     # 伏笔关联\n'
    '    emotionTag.ts         # 情绪标注\n'
    '    skeletonGen.ts        # 骨架方向生成\n'
    '  data/\n'
    '    templates.ts          # 5 个预设模板\n'
    '    emotionDict.ts        # 情绪关键词词典\n'
    '  types/\n'
    '    index.ts              # TS 类型定义\n'
    '  utils/\n'
    '    export.ts             # MD/TXT 导出\n'
    '    pwa.ts                # PWA 配置')

add_heading(doc, '1.3 核心数据模型（TypeScript 类型）', level=2)
add_mono(doc,
    'interface Card {\n'
    '  id: string\n'
    '  content: string            // <=500 字\n'
    '  createdAt: number\n'
    '  updatedAt: number\n'
    '  characterId?: string       // 关联角色\n'
    '  stage?: "pre"|"mid"|"post"|"none"  // 前期/中期/后期/未分类\n'
    '  emotion?: number          // -5..5\n'
    '  intensity?: number        // 1..5\n'
    '  order?: number            // 叙事线顺序\n'
    '}\n'
    '\n'
    'interface Character {\n'
    '  id: string\n'
    '  name: string\n'
    '  mentionCount: number\n'
    '  representativeDesc?: string\n'
    '  conflictTag?: string\n'
    '}\n'
    '\n'
    'interface Link {\n'
    '  id: string\n'
    '  cardAId: string\n'
    '  cardBId: string\n'
    '  reason: string            // 共现关键词\n'
    '  confirmed: boolean\n'
    '  hidden: boolean\n'
    '}\n'
    '\n'
    'interface Template {\n'
    '  id: string\n'
    '  name: string              // 追妻火葬场...\n'
    '  nodes: string[]           // 结构节点名\n'
    '  anchors: {name:string; hint:string}[]  // 关键锚点\n'
    '  idealEmotion: number[]    // 各节点理想情绪值\n'
    '}\n'
    '\n'
    'interface Chapter {\n'
    '  id: string\n'
    '  index: number\n'
    '  title: string\n'
    '  conflict: string          // 冲突一句话\n'
    '  cardIds: string[]         // 插入的卡片\n'
    '  nodeId: string            // 所属结构节点\n'
    '}')

add_heading(doc, '1.4 代码规范与工作流', level=2)
add_bullet(doc, '分支：main（生产）/ dev（开发）/ feat-xxx（功能分支）')
add_bullet(doc, '提交：feat: / fix: / refactor: / docs: / chore: 前缀')
add_bullet(doc, 'ESLint + Prettier 强制格式；Husky pre-commit 钩子')
add_bullet(doc, '每个任务一个 PR，至少自测通过再合并到 dev')

doc.add_page_break()

# ============== 2. 总体排期 ==============
add_heading(doc, '2. 总体排期与里程碑', level=1)

add_heading(doc, '2.1 Sprint 总览', level=2)
add_table(doc,
    ['Sprint', '主题', '覆盖需求', '任务数', '工时(人天)', '里程碑'],
    [
        ['S1', '项目搭建 + 阶段一卡片墙', 'D1, D2, A1, A3, A6', '9', '7', 'M1 阶段一可用'],
        ['S2', '阶段二梳理核心', 'B1, B2, B7', '7', '7', 'M2a 叙事线+角色'],
        ['S3', '阶段二梳理增强', 'B3, B4, B5, B6', '8', '7', 'M2b 梳理完整'],
        ['S4', '阶段三骨架+分区', 'C1, C2, C3', '7', '6', 'M3a 骨架生成'],
        ['S5', '阶段三章节+导出', 'C6, C7', '6', '5', 'M3b 大纲可导出'],
        ['S6', '情绪曲线+结构编辑', 'C4, C5', '6', '5', 'M3c 情绪可视化'],
        ['S7', '通用增强+双向流动', 'A2, A4, A5, C8, D3', '7', '5', 'M4 功能完整'],
        ['S8', '测试+优化+上线', '全量验收', '5', '5', 'M5 上线'],
    ],
    col_widths=[1.5, 4, 3.5, 1.5, 2, 3.5])
add_para(doc, '合计：8 个 Sprint，55 个任务，约 47 人天（按 1 人全职计）。', bold=True)

add_heading(doc, '2.2 里程碑时间线（按每周 5 工作日）', level=2)
add_table(doc,
    ['里程碑', '内容', '完成标志', '周次'],
    [
        ['M1', '阶段一可用', '能在网页新建/编辑/查看/删除卡片，数据持久化', '第 2 周末'],
        ['M2', '阶段二可用', '能一键梳理看叙事线/角色/伏笔/情绪/缺口', '第 4 周末'],
        ['M3', '阶段三可用', '能选模板生骨架、编辑章节、导出大纲文档', '第 6 周末'],
        ['M4', '功能完整', '情绪曲线+双向流动+数据导入导出', '第 7 周末'],
        ['M5', '上线', '测试通过，部署到 Vercel 可公开访问', '第 8 周末'],
    ],
    col_widths=[1.5, 5, 6.5, 3])

doc.add_page_break()

# ============== 3. Sprint 详细分解 ==============
add_heading(doc, '3. Sprint 详细任务分解', level=1)
add_para(doc, '每个任务格式：编号 | 任务 | 技术要点 | 产出 | 工时(h) | 依赖 | 验收', bold=True)

# ---- S1 ----
add_heading(doc, '3.1 Sprint 1：项目搭建 + 阶段一卡片墙', level=2)
add_para(doc, '目标：完成项目骨架，实现灵感卡片的增删改查与卡片墙展示，数据本地持久化。里程碑 M1。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T1.1', '项目初始化', 'Vite + React + TS；装依赖（Tailwind/Router/Zustand/Dexie/ECharts）；配 ESLint/Prettier/Husky', '可运行的空项目', '4', '-', 'pnpm dev 能启动空白页'],
        ['T1.2', '设计 token + 全局样式', 'Tailwind config 定义字号/间距/灰阶；配宋体；响应式断点 768/1024', 'tailwind.config + globals.css', '3', 'T1.1', '基础样式生效'],
        ['T1.3', '路由 + 全局布局', 'React Router 配 4 路由；Navbar 组件（桌面/移动形态切换）；阶段进度条', 'App.tsx + Navbar.tsx', '4', 'T1.2', '4 路由可切换'],
        ['T1.4', 'IndexedDB 数据层', 'Dexie 定义 cards/characters/links/chapters 表；cardRepo 实现 CRUD', 'db/database.ts + cardRepo.ts', '5', 'T1.1', 'CRUD 单测通过'],
        ['T1.5', '类型定义', 'types/index.ts 定义 Card/Character/Link/Template/Chapter', 'types/index.ts', '2', 'T1.1', '类型可被引用'],
        ['T1.6', 'cardStore 状态管理', 'Zustand store：cards 列表 + loadAll/add/update/remove；订阅 Dexie', 'stores/cardStore.ts', '4', 'T1.4, T1.5', 'store 与 DB 同步'],
        ['T1.7', '卡片墙页面布局', 'CSS Grid auto-fill minmax(260px)；时间倒序；移动端单列', 'routes/CardsPage.tsx', '5', 'T1.3, T1.6', '卡片正确排列'],
        ['T1.8', '卡片缩略卡组件', 'CardThumb：缩略2行/移动4行、角色标签、阶段标签、时间戳；悬停上移', 'components/CardThumb.tsx', '4', 'T1.7', '缩略卡正确渲染'],
        ['T1.9', '新建/编辑弹窗', 'CardEditModal：textarea 500字限制+计数、角色下拉、阶段单选、保存/取消', 'components/CardEditModal.tsx', '6', 'T1.8', '能新建+编辑卡片'],
        ['T1.10', '卡片删除+空态', '详情弹窗删除二次确认；0 张时空态插图+引导按钮', 'CardDetailModal + EmptyState', '3', 'T1.9', '删除+空态正常'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S1 合计 40h（约 7 人天）。验收：能在 /cards 完成卡片增删改查，刷新数据不丢，手机电脑布局正常。', bold=True)

# ---- S2 ----
add_heading(doc, '3.2 Sprint 2：阶段二梳理核心（叙事线 + 角色）', level=2)
add_para(doc, '目标：实现「一键梳理」的叙事线排列和角色提取，产出可查看的梳理结果。里程碑 M2a。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T2.1', 'jieba-wasm 集成', '装 jieba-wasm；封装 tokenize(text) 返回词数组；加载词典', 'utils/tokenizer.ts', '5', 'T1.5', '能分词中文'],
        ['T2.2', '叙事线排序引擎', '权重：stage(前期3/中期2/后期1/未分类0) > 时间戳 > 关键词(开始/初见/结局/最后)；输出 order', 'engine/narrativeLine.ts', '6', 'T2.1', '排序结果合理'],
        ['T2.3', '缺口检测', '按 order 计算相邻卡片间距；超阈值标缺口 + 提示文案', 'engine/narrativeLine.ts', '3', 'T2.2', '缺口正确标注'],
        ['T2.4', '角色提取引擎', 'jieba 分词 + 人名词典过滤；统计频率；取代表性描述（出现最多的卡片原文）', 'engine/characterExtract.ts', '5', 'T2.1', '角色列表正确'],
        ['T2.5', '冲突标签生成', '规则：角色在不同卡片情绪正负矛盾时标"嘴上说放下，卡片里反复写他"', 'engine/characterExtract.ts', '3', 'T2.4', '标签合理'],
        ['T2.6', '梳理结果页骨架', '/sort 路由；顶部标签页（叙事线/角色/关联）；梳理中 loading 动画', 'routes/SortPage.tsx', '4', 'T1.3', '页面可进入'],
        ['T2.7', '叙事线视图', '纵向时间线渲染卡片 + 缺口虚线；可拖动改 order', 'components/TimelineView.tsx', '5', 'T2.2, T2.6', '时间线正确'],
        ['T2.8', '角色视图', '角色列表（频率排序）+ 代表性描述 + 冲突标签', 'components/CharacterList.tsx', '3', 'T2.4, T2.6', '角色列表正确'],
        ['T2.9', '进入大纲构建入口', '"进入大纲构建"按钮跳 /outline；传已确认叙事线+角色', 'SortPage 按钮', '1', 'T2.7', '能跳转阶段三'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S2 合计 35h（约 7 人天）。验收：点梳理后 15 秒内出叙事线+角色，缺口有提示，可拖动卡片。', bold=True)

# ---- S3 ----
add_heading(doc, '3.3 Sprint 3：阶段二梳理增强（伏笔 + 缺口 + 情绪 + 操作）', level=2)
add_para(doc, '目标：补全梳理的伏笔关联、情绪标注、结果操作。里程碑 M2b。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T3.1', '伏笔关联引擎', '两两卡片关键词共现统计；共现≥2 标关联；生成 reason 文案', 'engine/foreshadowLink.ts', '5', 'T2.1', '关联对合理'],
        ['T3.2', '关联视图', '关联标签页；卡片对并排展示 + 系统标注；点进看全文', 'components/LinkView.tsx', '4', 'T3.1, T2.6', '关联可查看'],
        ['T3.3', '关联操作', '标记已确认（虚线→实线）/ 删除（隐藏）；存 DB', 'LinkView + linkRepo', '3', 'T3.2', '操作持久化'],
        ['T3.4', '情绪关键词词典', '约 200 词，分正负+强度；data/emotionDict.ts', 'data/emotionDict.ts', '3', '-', '词典可用'],
        ['T3.5', '情绪标注引擎', '扫描卡片内容匹配词典；计算默认 emotion/intensity', 'engine/emotionTag.ts', '4', 'T3.4', '默认值合理'],
        ['T3.6', '卡片情绪编辑', '卡片详情弹窗加情绪滑块(-5..5)+强度(1..5)；保存到 DB', 'CardDetailModal 扩展', '3', 'T3.5, T1.9', '可手动改情绪'],
        ['T3.7', '缺口标注细化', '密集处提示"考虑拆分"；稀疏处提示"需要转折"', 'narrativeLine.ts', '2', 'T2.3', '提示文案正确'],
        ['T3.8', '梳理结果缓存', '同批卡片未改则返回上次结果；cardIds hash 做缓存键', 'sortStore', '3', 'T2.2', '二次梳理秒出'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S3 合计 27h（约 5.5 人天）。验收：伏笔关联可标可删，情绪有默认值可改，缺口提示细化。', bold=True)

# ---- S4 ----
add_heading(doc, '3.4 Sprint 4：阶段三骨架 + 分区编辑器', level=2)
add_para(doc, '目标：实现 5 个预设模板、生成 2-3 条骨架方向、时间线分区编辑器。里程碑 M3a。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T4.1', '5 个预设模板数据', '追妻火葬场/破镜重圆/替身白月光/先婚后爱/重生复仇；含 nodes/anchors/idealEmotion', 'data/templates.ts', '4', 'T1.5', '模板数据完整'],
        ['T4.2', '骨架方向生成引擎', '不同模板或同模板不同初始分区方案生成 2-3 条；每条含分区边界+锚点初始位置', 'engine/skeletonGen.ts', '6', 'T4.1, T2.2', '能生成多方向'],
        ['T4.3', '/outline 页面骨架', '左侧模板栏 + 中部叙事线+分区 + 右侧章节区', 'routes/OutlinePage.tsx', '4', 'T1.3', '页面可进入'],
        ['T4.4', '模板选择 + 方向对比', '选模板后展示 2-3 条方向卡片；每条预览分区+锚点', 'components/DirectionPicker.tsx', '4', 'T4.2, T4.3', '能选方向'],
        ['T4.5', '时间线分区编辑器', '叙事线上叠加分区条；分界线默认均分可拖动；锚点可拖到卡片', 'components/TimelineDivider.tsx', '8', 'T4.4', '分区可拖动'],
        ['T4.6', '卡片归段', '卡片落在哪段归该段章节池；分界线移动时自动重算', 'outlineStore', '3', 'T4.5', '归段正确'],
        ['T4.7', '回阶段一补灵感', '"回阶段一"按钮跳 /cards；补完后回 /outline 刷新', 'OutlinePage 按钮', '2', 'T4.3', '能来回跳'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S4 合计 31h（约 6 人天）。验收：能选模板看 2-3 条方向，选一条后分区线可拖动，锚点可锚定卡片。', bold=True)

# ---- S5 ----
add_heading(doc, '3.5 Sprint 5：阶段三章节编辑 + 大纲导出', level=2)
add_para(doc, '目标：实现章节拆合、卡片入章、冲突编辑、大纲文档导出。里程碑 M3b。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T5.1', '章节自动拆分', '按段内卡片数建议（3内1章/4-6两章/7+三章）；生成 Chapter 列表', 'skeletonGen.ts', '3', 'T4.6', '拆分合理'],
        ['T5.2', '章节列表组件', '每章一行：序号+标题+冲突+卡片编号；右侧卡片插入槽', 'components/ChapterList.tsx', '5', 'T5.1', '列表正确渲染'],
        ['T5.3', '章节编辑交互', '拆合章、拖卡片入章、双击改冲突描述、增删章节', 'ChapterList + outlineStore', '6', 'T5.2', '编辑可持久化'],
        ['T5.4', '大纲文档生成', '组装 MD/TXT：书名+模板+逐章+卡片引用+角色弧光+伏笔清单+情绪说明', 'utils/export.ts', '5', 'T5.3', '文档内容完整'],
        ['T5.5', '/export 导出页', '文档预览区 + 导出按钮（MD/TXT）；file-saver 下载', 'routes/ExportPage.tsx', '4', 'T5.4', '能下载文件'],
        ['T5.6', '回阶段二重新确认', '"回阶段二"按钮跳 /sort；不覆盖阶段三已有编辑', 'OutlinePage 按钮', '2', 'T5.3', '不覆盖编辑'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S5 合计 25h（约 5 人天）。验收：能编辑章节，导出的 MD 文件含全部要素。', bold=True)

# ---- S6 ----
add_heading(doc, '3.6 Sprint 6：情绪曲线 + 结构编辑器', level=2)
add_para(doc, '目标：实现情绪冲突曲线可视化和结构节点自定义。里程碑 M3c。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T6.1', 'ECharts 情绪曲线', '横轴卡片顺序，纵轴情绪值；实线+圆点（大小=强度）', 'components/EmotionChart.tsx', '5', 'T3.5, T4.5', '曲线正确渲染'],
        ['T6.2', '模板参考线叠加', '半透明虚线显示当前模板 idealEmotion', 'EmotionChart 扩展', '3', 'T6.1', '参考线显示'],
        ['T6.3', '分区竖线叠加', '分区边界在曲线上画竖虚线', 'EmotionChart 扩展', '2', 'T6.1', '竖线对齐分区'],
        ['T6.4', '曲线节点拖动', '拖圆点上下改情绪值；点击圆点跳转卡片', 'EmotionChart + cardStore', '4', 'T6.1', '拖动可改值'],
        ['T6.5', '结构编辑器', '节点改名/增删(2-6个)/调顺序；更新模板 nodes', 'components/StructureEditor.tsx', '5', 'T4.1', '能自定义结构'],
        ['T6.6', '结构变更联动', '改节点后分区条+章节+曲线自动跟随', 'outlineStore', '3', 'T6.5', '联动正确'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S6 合计 22h（约 4.5 人天）。验收：情绪曲线可拖改，结构节点可自定义并联动。', bold=True)

# ---- S7 ----
add_heading(doc, '3.7 Sprint 7：通用增强 + 双向流动', level=2)
add_para(doc, '目标：补全 P1 增强：卡片元数据、角色筛选、轻提示、数据导入导出。里程碑 M4。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T7.1', '卡片元数据补全', '新建弹窗加角色选择(复用已有)+阶段选择', 'CardEditModal 扩展', '3', 'T1.9', '元数据可填'],
        ['T7.2', '角色筛选', '顶栏角色标签点击筛选；再点取消；"全部"重置', 'CardsPage 筛选栏', '3', 'T7.1', '筛选生效'],
        ['T7.3', '卡片少时轻提示', '<5 张点梳理弹确认对话框', 'CardsPage 提示', '1', 'T2.6', '提示正确触发'],
        ['T7.4', '数据导出 JSON', '全量 cards/characters/links/chapters 导出 .json', 'utils/backup.ts', '3', 'T1.4', '能下载备份'],
        ['T7.5', '数据导入 JSON', '文件选择 + 解析 + 校验 + 写入 DB', 'utils/backup.ts', '3', 'T7.4', '能恢复数据'],
        ['T7.6', 'PWA 配置', 'vite-plugin-pwa；manifest + service worker；可离线', 'vite.config + pwa.ts', '3', 'T1.1', '离线可用'],
        ['T7.7', '回阶段二不覆盖逻辑', '阶段三调整后回阶段二，标记"已有阶段三编辑"，确认才覆盖', 'outlineStore', '2', 'T5.6', '不覆盖编辑'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S7 合计 18h（约 3.5 人天）。验收：筛选/提示/导入导出/PWA 全部可用。', bold=True)

# ---- S8 ----
add_heading(doc, '3.8 Sprint 8：测试 + 优化 + 上线', level=2)
add_para(doc, '目标：全量测试、性能优化、部署上线。里程碑 M5。', bold=True)
add_table(doc,
    ['编号', '任务', '技术要点', '产出', '工时(h)', '依赖', '验收'],
    [
        ['T8.1', '功能测试', '逐项验收 P0 需求；记录 bug', '测试报告', '8', 'S1-S7', 'P0 全通过'],
        ['T8.2', '性能优化', '100 张卡片场景：虚拟列表、梳理 Web Worker、曲线节流', '代码优化', '6', 'T8.1', '100 卡 <15s'],
        ['T8.3', '兼容测试', 'Chrome/Safari/Edge/iOS/Android 实测', '兼容报告', '4', 'T8.1', '主流浏览器正常'],
        ['T8.4', '移动端交互优化', '拖动改长按+点击（若体验差）；按钮可达性', '交互调整', '4', 'T8.1', '移动端可用'],
        ['T8.5', '部署上线', 'Vercel 部署；域名绑定；构建优化', '线上站点', '3', 'T8.2', '可公开访问'],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3, 1.5, 3])
add_para(doc, 'S8 合计 25h（约 5 人天）。验收：测试通过，线上可访问。', bold=True)

doc.add_page_break()

# ============== 4. 依赖关系 ==============
add_heading(doc, '4. 任务依赖关系', level=1)
add_para(doc, '关键依赖链（决定项目最长路径）：', bold=True)
add_mono(doc,
    'T1.1 初始化\n'
    '  -> T1.4 数据层 -> T1.6 store -> T1.7 卡片墙 -> T1.9 弹窗 (M1)\n'
    '  -> T2.1 分词 -> T2.2 叙事线 -> T2.7 时间线视图 (M2a)\n'
    '           -> T2.4 角色 -> T2.8 角色视图\n'
    '           -> T3.1 伏笔 -> T3.2 关联视图\n'
    '           -> T3.5 情绪标注 -> T6.1 情绪曲线\n'
    '  -> T4.1 模板 -> T4.2 骨架生成 -> T4.5 分区编辑器 (M3a)\n'
    '                              -> T5.1 章节拆分 -> T5.3 章节编辑 -> T5.4 导出 (M3b)\n'
    '  -> T6.5 结构编辑器 (M3c)\n'
    '  -> T7.* 增强 (M4)\n'
    '  -> T8.* 测试上线 (M5)')
add_para(doc, '')
add_para(doc, '可并行任务（无依赖，可同时推进）：', bold=True)
add_bullet(doc, 'T3.4 情绪词典（独立）与 T2.* 梳理引擎并行')
add_bullet(doc, 'T4.1 模板数据（独立）与 T3.* 并行')
add_bullet(doc, 'T7.6 PWA 配置 与 S5/S6 并行')

doc.add_page_break()

# ============== 5. 技术风险 ==============
add_heading(doc, '5. 技术风险与应对', level=1)
add_table(doc,
    ['编号', '风险', '等级', '应对措施', '验证节点'],
    [
        ['TR1', 'jieba-wasm 体积大/加载慢', '高', 'S1 末做体积测试；备选：预编译精简词典或用 nodejieba 预处理', 'T2.1 完成时'],
        ['TR2', 'ECharts 拖动节点性能差', '中', 'S6 初做 POC；备选：自研 SVG 曲线', 'T6.1 完成时'],
        ['TR3', '移动端拖动分界线体验差', '中', 'S8 兼容测试验证；备选：长按+点击+输入框精确调整', 'T8.4'],
        ['TR4', 'IndexedDB 在 iOS Safari 配额限制', '中', 'S8 测试；提供导出备份；数据量预估 < 5MB', 'T8.3'],
        ['TR5', '情绪词典覆盖不全导致默认值不准', '低', 'S3 上线后迭代；提供"重扫所有卡片"功能', 'M2b 后'],
        ['TR6', '叙事线排序规则效果不佳', '中', 'S2 用真实卡片数据测试；权重可调', 'T2.2 完成时'],
    ],
    col_widths=[1.2, 4.5, 1.5, 5.5, 3.3])

doc.add_page_break()

# ============== 6. 验收检查清单 ==============
add_heading(doc, '6. 验收检查清单', level=1)
add_para(doc, '按 P0 需求逐项验收，每项打勾才算通过。', bold=True)

add_heading(doc, '6.1 阶段一', level=2)
add_table(doc,
    ['需求', '验收项', '通过'],
    [
        ['A1', '能新建文字卡片（<=500字）+ 时间戳', '□'],
        ['A3', '卡片墙时间倒序 + 缩略前2行', '□'],
        ['A6', '卡片可双击编辑 + 可删除（二次确认）', '□'],
        ['D1', '手机 375px / 电脑 1440px 布局正常', '□'],
        ['D2', '刷新数据不丢失（IndexedDB 持久化）', '□'],
    ],
    col_widths=[1.5, 12, 2.5])

add_heading(doc, '6.2 阶段二', level=2)
add_table(doc,
    ['需求', '验收项', '通过'],
    [
        ['B1', '梳理后卡片按故事顺序排成时间线 + 缺口提示', '□'],
        ['B2', '角色按频率排序 + 代表性描述 + 冲突标签', '□'],
        ['B3', '伏笔关联成对展示 + 系统标注', '□'],
        ['B5', '情绪值有默认值 + 可手动改', '□'],
        ['B6', '可标记已确认 / 删除关联 / 拖动卡片', '□'],
        ['B7', '能跳转阶段三', '□'],
    ],
    col_widths=[1.5, 12, 2.5])

add_heading(doc, '6.3 阶段三', level=2)
add_table(doc,
    ['需求', '验收项', '通过'],
    [
        ['C1', '5 个预设模板可选', '□'],
        ['C2', '生成 2-3 条骨架方向', '□'],
        ['C3', '分区线可拖动 + 锚点可锚定卡片', '□'],
        ['C4', '情绪曲线显示 + 模板参考线 + 可拖圆点', '□'],
        ['C6', '章节可拆合 / 拖卡 / 改冲突 / 增删', '□'],
        ['C7', '导出 MD/TXT 含角色弧光+伏笔清单+情绪说明', '□'],
    ],
    col_widths=[1.5, 12, 2.5])

add_heading(doc, '6.4 通用', level=2)
add_table(doc,
    ['需求', '验收项', '通过'],
    [
        ['A2', '卡片可填关联角色 + 所属阶段', '□'],
        ['A4', '角色筛选生效', '□'],
        ['A5', '<5 张卡片时轻提示', '□'],
        ['C8', '阶段间双向流动 + 不覆盖编辑', '□'],
        ['D3', 'JSON 导入导出可用', '□'],
        ['PWA', '可离线使用', '□'],
    ],
    col_widths=[1.5, 12, 2.5])

doc.add_page_break()

# ============== 7. 开发环境准备 ==============
add_heading(doc, '7. 开发环境准备清单', level=1)
add_para(doc, '开工前需准备好的环境：', bold=True)
add_table(doc,
    ['项目', '要求', '状态'],
    [
        ['Node.js', '>= 18 LTS', '□'],
        ['包管理', 'pnpm（推荐）或 npm', '□'],
        ['编辑器', 'VS Code + ESLint/Prettier 插件', '□'],
        ['Git', '已配置；GitHub 仓库已建', '□'],
        ['Vercel 账号', '用于部署', '□'],
        ['测试设备', 'Chrome + 手机浏览器（iOS/Android）', '□'],
        ['真实数据', '准备 20-30 条女频灵感卡片用于测试', '□'],
    ],
    col_widths=[3, 9, 4])

# ============== 保存 ==============
output_path = r'd:\app\小说大纲辅助器\小说大纲梳理器开发计划_v1.0.docx'
doc.save(output_path)
print(f'开发计划已生成：{output_path}')
